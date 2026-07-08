from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT_DIR = BASE_DIR / "docs"

COLOR_H1 = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H2 = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H3 = RGBColor(0x1F, 0x4D, 0x78)
COLOR_TEXT = RGBColor(0x19, 0x35, 0x1D)
COLOR_MUTED = RGBColor(0x54, 0x70, 0x5A)
COLOR_TABLE_FILL = RGBColor(0xE8, 0xEE, 0xF5)


def set_document_base(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_TEXT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for style_name, size, color, before, after, bold in (
        ("Heading 1", 16, COLOR_H1, 18, 10, True),
        ("Heading 2", 13, COLOR_H2, 14, 7, True),
        ("Heading 3", 12, COLOR_H3, 10, 5, True),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color_hex)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if row_index == 0:
            for cell in row.cells:
                set_cell_shading(cell, "E8EEF5")
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def add_title_block(document: Document, title: str, subtitle: str) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(8)
    title_run = title_paragraph.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_H1

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.paragraph_format.space_before = Pt(0)
    subtitle_paragraph.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = COLOR_MUTED


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.19)
    run = paragraph.add_run(f"• {text}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_numbered(document: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.left_indent = Inches(0.38)
        paragraph.paragraph_format.first_line_indent = Inches(-0.19)
        run = paragraph.add_run(f"{index}. {item}")
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def build_scheme_doc() -> Path:
    document = Document()
    set_document_base(document)
    add_title_block(
        document,
        "Esquema de Trabajo del Portal de Horarios",
        "Flujo operativo para el cargue de usuarios por sede y el uso del perfil administrador para consulta y control.",
    )

    intro = document.add_paragraph(
        "Este esquema resume como se registra un usuario, como se le asigna su alcance por sede y como opera el perfil administrador dentro del portal."
    )
    intro.paragraph_format.space_after = Pt(10)

    document.add_paragraph("1. Flujo General", style="Heading 1")
    add_numbered(
        document,
        [
            "El administrador ingresa al modulo de administracion de Django.",
            "Crea el usuario y guarda el registro base.",
            "Edita el usuario y define el acceso al portal.",
            "Si el perfil es Administrador, el usuario podra operar todas las sedes.",
            "Si el perfil es Usuario por sede, solo podra ver y editar las sedes asignadas.",
            "El usuario inicia sesion y el portal muestra las opciones permitidas segun su perfil.",
        ],
    )

    document.add_paragraph("2. Cargue Usuario-Sede", style="Heading 1")
    table = document.add_table(rows=1, cols=4)
    table.columns[0].width = Inches(0.9)
    table.columns[1].width = Inches(1.9)
    table.columns[2].width = Inches(1.8)
    table.columns[3].width = Inches(1.9)
    headers = ["Paso", "Accion", "Responsable", "Resultado"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    rows = [
        ("1", "Crear usuario en Django Admin.", "Administrador", "Usuario base creado."),
        ("2", "Ingresar a la edicion del usuario.", "Administrador", "Se habilita el acceso al portal."),
        ("3", "Definir perfil Administrador o Usuario por sede.", "Administrador", "Queda configurado el alcance del usuario."),
        ("4", "Asignar una o varias sedes cuando el perfil no es administrador.", "Administrador", "El usuario queda restringido a esas sedes."),
        ("5", "Guardar cambios y validar acceso con inicio de sesion.", "Administrador", "Usuario listo para operar."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table)

    document.add_paragraph("3. Uso del Usuario por Sede", style="Heading 1")
    add_bullet(document, "Ingresa al portal con su usuario y contrasena.")
    add_bullet(document, "Solo visualiza las sedes que tiene asignadas.")
    add_bullet(document, "Puede crear, editar y diligenciar horarios de esas sedes.")
    add_bullet(document, "Puede consultar personal cargado, turnos, pendientes y validaciones del horario.")
    add_bullet(document, "No puede eliminar horarios ni administrar parametros globales si no tiene perfil administrador.")

    document.add_paragraph("4. Uso del Administrador para Consulta", style="Heading 1")
    add_bullet(document, "Consulta todas las sedes activas del sistema.")
    add_bullet(document, "Revisa el dashboard general con indicadores de estado, borradores, publicados y alertas.")
    add_bullet(document, "Abre cualquier horario por sede y semana para revisar totales, extras, recargos, pendientes, pagos y alertas.")
    add_bullet(document, "Administra sedes, cargos, turnos y configuracion general del sistema.")
    add_bullet(document, "Controla accesos por usuario y puede eliminar horarios cuando el proceso lo requiera.")

    document.add_paragraph("5. Reglas de Operacion", style="Heading 1")
    add_bullet(document, "Administrador: ve todo, configura todo y consulta toda la operacion.")
    add_bullet(document, "Usuario por sede: solo trabaja sobre la sede o sedes asignadas.")
    add_bullet(document, "Los horarios publicados quedan cerrados para edicion y se consultan en modo lectura.")
    add_bullet(document, "Los datos de horas, extras, recargos, dias pendientes y pagos quedan almacenados en la base de datos del portal.")

    output_path = OUTPUT_DIR / "Esquema_Trabajo_Portal_Horarios.docx"
    document.save(output_path)
    return output_path


def build_manual_doc() -> Path:
    document = Document()
    set_document_base(document)
    add_title_block(
        document,
        "Manual Operativo del Portal de Horarios",
        "Guia practica para administradores y usuarios responsables del diligenciamiento, consulta y control del horario semanal.",
    )

    document.add_paragraph("1. Objetivo", style="Heading 1")
    document.add_paragraph(
        "Establecer el uso operativo del portal para la programacion semanal de personal por sede, el control de turnos y la consulta de indicadores asociados al horario."
    )

    document.add_paragraph("2. Perfiles del Sistema", style="Heading 1")
    table = document.add_table(rows=1, cols=3)
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(2.2)
    table.columns[2].width = Inches(2.8)
    for idx, header in enumerate(["Perfil", "Alcance", "Funciones principales"]):
        table.cell(0, idx).text = header
    profile_rows = [
        ("Administrador", "Todas las sedes", "Configura catalogos, consulta dashboard, revisa horarios, administra usuarios y elimina horarios."),
        ("Usuario por sede", "Sedes asignadas", "Crea horarios, diligencia turnos, registra pendientes y consulta solo la informacion permitida."),
    ]
    for row in profile_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table)

    document.add_paragraph("3. Ingreso al Portal", style="Heading 1")
    add_numbered(
        document,
        [
            "Abrir la URL interna del portal.",
            "Ingresar usuario y contrasena.",
            "Validar que el menu lateral muestre las opciones acordes al perfil.",
            "Si el acceso falla, solicitar revision de credenciales o perfil asignado.",
        ],
    )

    document.add_paragraph("4. Procedimiento para el Administrador", style="Heading 1")
    document.add_paragraph("4.1 Gestion de usuarios y sedes", style="Heading 2")
    add_numbered(
        document,
        [
            "Entrar al admin de Django.",
            "Crear o editar el usuario.",
            "Definir si el perfil sera Administrador o Usuario por sede.",
            "Asignar las sedes correspondientes cuando aplique.",
            "Guardar y validar el acceso en el portal.",
        ],
    )

    document.add_paragraph("4.2 Consulta general de la operacion", style="Heading 2")
    add_bullet(document, "Revisar el dashboard para identificar borradores, horarios publicados y alertas.")
    add_bullet(document, "Entrar a Sedes para confirmar catalogo disponible y accesible.")
    add_bullet(document, "Entrar a Horarios para abrir semanas ya creadas o generar nuevas.")
    add_bullet(document, "Validar totales de horas, horas extra, recargo nocturno, dias pendientes, horas pendientes, pagos y alertas.")

    document.add_paragraph("5. Procedimiento para el Usuario por Sede", style="Heading 1")
    add_numbered(
        document,
        [
            "Ingresar al portal y seleccionar la sede permitida.",
            "Definir el dia inicio de semana y abrir la grilla de horario.",
            "Confirmar que el personal cargado corresponda a la sede.",
            "Asignar turnos diarios por trabajador usando las listas desplegables.",
            "Registrar novedades, pendientes, dias por pagar y horas por pagar cuando aplique.",
            "Guardar el horario y revisar las alertas antes de dejarlo listo.",
        ],
    )

    document.add_paragraph("6. Diligenciamiento del Horario", style="Heading 1")
    add_bullet(document, "Turno 1 y Turno 2 deben guardar coherencia de jornada.")
    add_bullet(document, "Si el trabajador cumple la jornada en un solo turno, no debe llevar segundo turno.")
    add_bullet(document, "Los calculos de horas del dia, total semanal, extras y recargos se actualizan automaticamente.")
    add_bullet(document, "Las fechas pendientes deben coincidir con la cantidad de dias pendientes.")
    add_bullet(document, "Los pagos por dia y por horas solo deben aplicarse cuando exista saldo previo acumulado.")

    document.add_paragraph("7. Estados del Horario", style="Heading 1")
    state_table = document.add_table(rows=1, cols=3)
    state_table.columns[0].width = Inches(1.4)
    state_table.columns[1].width = Inches(2.0)
    state_table.columns[2].width = Inches(3.1)
    for idx, header in enumerate(["Estado", "Uso", "Comportamiento"]):
        state_table.cell(0, idx).text = header
    state_rows = [
        ("Borrador", "Horario en construccion", "Permite edicion completa y ajustes del personal."),
        ("En revision", "Validacion previa", "Se consulta y se corrige antes de publicarlo."),
        ("Publicado", "Horario cerrado", "Queda en modo lectura y ya no permite modificaciones."),
    ]
    for row in state_rows:
        cells = state_table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(state_table)

    document.add_paragraph("8. Indicadores que Deben Revisarse", style="Heading 1")
    add_bullet(document, "Horas totales de la semana por trabajador.")
    add_bullet(document, "Horas extra generadas frente a la meta semanal del cargo.")
    add_bullet(document, "Recargo nocturno cuando existan turnos desde las 7:00 p. m. en adelante.")
    add_bullet(document, "Dias pendientes y horas pendientes acumuladas.")
    add_bullet(document, "Pagos aplicados por dia o por horas.")
    add_bullet(document, "Delta y alertas para identificar diferencias o excesos.")

    document.add_paragraph("9. Buenas Practicas Operativas", style="Heading 1")
    add_bullet(document, "Guardar cambios con frecuencia durante el diligenciamiento.")
    add_bullet(document, "Validar alertas antes de publicar el horario.")
    add_bullet(document, "No usar turnos o novedades que no correspondan al caso real.")
    add_bullet(document, "Mantener actualizados los catalogos de cargos, turnos y sedes.")
    add_bullet(document, "Usar comentarios o notas del horario cuando se requiera dejar contexto adicional.")

    document.add_paragraph("10. Soporte y Escalamiento", style="Heading 1")
    document.add_paragraph(
        "Si el sistema presenta errores de acceso, no carga personal, no guarda el horario o genera validaciones que no correspondan, el caso debe escalarse al administrador funcional o al equipo tecnico responsable del portal para revision de configuracion, datos o despliegue."
    )

    output_path = OUTPUT_DIR / "Manual_Operativo_Portal_Horarios.docx"
    document.save(output_path)
    return output_path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_scheme_doc()
    build_manual_doc()


if __name__ == "__main__":
    main()
